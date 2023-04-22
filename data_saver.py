import os
from typing import List, Dict
from loguru import logger
from blast_ncbi import BlastNCBIResults

import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX


def add_hyperlink(paragraph, url, text):

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(
        docx.oxml.shared.qn("r:id"),
        r_id,
    )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement("w:r")
    rPr = docx.oxml.shared.OxmlElement("w:rPr")

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def save_results_in_word(path: str, file_name: str, species: List[BlastNCBIResults]):
    # Extract names to be added to the Word document table columns
    column_names = [
        name.replace("_", " ").capitalize()
        for name in species[0].__annotations__.keys()
        if "url" not in name
    ]

    # Create an instance of a word document
    doc = docx.Document()

    # Add a Title to the document
    doc.add_heading("Sequences producing significant alignments", 0)

    # Creating a table object
    table = doc.add_table(rows=1, cols=len(column_names), style="Table Grid")

    # Adding heading in the 1st row of the table
    for cell, column_name in zip(table.rows[0].cells, column_names):
        cell.text = column_name

    attr_names = [name for name in species[0].__annotations__.keys() if "url" not in name]
    for specie in species:

        for attribute, cell in zip(attr_names, table.add_row().cells):

            # Try to add hyperlink
            try:
                href = getattr(specie, f"{attribute}_url")
                p_table = cell.add_paragraph()
                add_hyperlink(p_table, href, getattr(specie, attribute))
                exist_link = True
            except AttributeError:
                exist_link = False

            # If the attribute has no _url, then simply add the text to the table cell
            if not exist_link:
                cell.text = getattr(specie, attribute)

    # Create directories in disk memory
    if not os.path.exists(path):
        os.makedirs(path)

    # Now save the document to a location
    doc.save(f"{path}/{file_name}.docx")


def save_alignments_to_notes(path: str, file_name: str, alignments: Dict[str, str]):

    # Create directories in disk memory
    if not os.path.exists(path):
        os.makedirs(path)

    with open(f"{path}/{file_name}.txt", "w") as file:
        for id, sequence in alignments.items():
            file.write(f"{sequence}\n")


if __name__ == "__main__":
    alignments = {
        0: "GGCACTGCGGC-TGCCTATAC-TGCAAGTTCGAGCGAATGGATTGAGAAGCTTGCTTCTCAAGAAGTTAGCGGCGGACGGGTGAGTAACACGTGGGTAACCTGCCCATAAGAGTGGGATAACTCCGGGAAACCGGGGCTAATACCGGATAATATTTTGAACTGCATGGTTCGAAATTGAAAGGCGGCTTCGGCTGTCACTTATGGATGGACCCGCGTCGCATTAGCTAGTTGGTGAGGTAACGGCTCACCAAGGCAACGATGCGTAGCCGACCTGAGAGGGTGATCGGCCACACTGGGACTGAGACACGGCCCAGACTCCTACGGGAGGCAGCAGTAGGGAATCTTCCGCAATGGACGAAAGTCTGACGGAGCAACGCCGCGTGAGTGATGAAGGCTTTCGGGTCGTAAAACTCTGTTGTTAGGGAAGAACAAGTGCTAGTTGAATAAGCTGGCACCTTGACGGTACCTAACCAGAAAGCCACGGCTAACTACGTGCCAGCAGCCGCGGTAATACGTAGGTGGCAAGCGTTATCCGGAATTATTGGGCGTAAAGCGCGCGCAGGTGGTTTCTTAAGTCTGATGTGAAAGCCCACGGCTCAACCGTGGAGGGTCATTGGAAACTGGGAGACTTGAGTGCAGAAGAGGAAAGTGGAATTCCTGGTGTAGCGGTGAAATGCGTAGAGATATGGAGGAACACCAGTGGCGAAAGCGACTTTCTGGTCTGTAACTGACACTGAGGCGCGAAAGCGTGGGGAGCAAACAGGATTAGATACCCTGGTAGTCCACGCCGTAAACGATGAGTGCTAAGTGTTTAAAGGGTTTCCGCCCTTTAGTGCTGAAGTTAACGCATTAAGCACTCCGCCCGGGGGAGTACGGCCGCAAGGCTGAAACTCAAAGGAATTGACGGGGGCCCGCACAAGCGGTGGAGCATGTGGTTTAATTCGAAGCAACGCGAAGAACCTTACCAGGTCTTGACATCCTCTGAAAACCCTAGAGATAGGGCTTCTCCTTCGGGAGCAGAATGACAGGTGGTGCAAGGGTTGTCTTCCCCTCCTGTCCTGAGATATTTGGGTTTATTTCCTCCACCGA-CGCCACCCCTTGTTCT-ATTTTCT-TCCTTAATTTGGGC  1125",
        1: "              .......-...-..-..................-........C.......................................................C..................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................AT...............................................G.....................................................................................................TG...G................................................T-...........................................................................................................................................................................T........G..A...............-.G.....G...AG...C...A.A.G...A....-...A...                         1092",
        2: "....T..GC..-A........A...-..-...................-......-........................................................C......................................C...........................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................AT...............................................G......................................................................................................-...G..............................................-.................................................................A.................................................................A.A....................A..A.................-.......C..A............-..A....T..G...A..........A.-...A....-...A...T....G.CC......          1111",
        3: "      .....-.........A...-..................A..-...........T.T..................................................C......................................C.........C.................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................AT...............................................G......................................................................................................-...G..............................................-...........................................................................................................................C.................................A...............-.........C..A.............A.AT........A... ..G.A..-..G...A....-...A...T.G..G.CA..A....G......  1127",
        4: "      .....G.........A......-...................-......-........................................................C..................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................AT...............................................G......................................................................................................-...G................................................T-............................................................................................................................................................G..............T-.......G..AG...G....G......G-.......A.G.C..G.A..-..G...A....-...A...T.G..G.CA..A....G......  1127",
        5: "      .....G.........A......-...................-......-........................................................C..................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................AT...............................................G......................................................................................................-...G................................................T-............................................................................................................................................................G..............T-.......G..AG...G....G......G-.......A.G.C",
    }
    species = [
        BlastNCBIResults(
            description="Bacillus toyonensis strain SL4-3 16S ribosomal RNA gene partial sequence",
            description_url="https://blast.ncbi.nlm.nih.gov/Blast.cgi#alnHdr_1538993297",
            scientific_name="Bacillus toyonensis",
            scientific_name_url="https://www.ncbi.nlm.nih.gov/Taxonomy/Browser/wwwtax.cgi?id=155322",
            max_score="1842",
            total_score="1842",
            query_cover="80%",
            e_value="0.0",
            per_indentity="96.44%",
            accession_len="1470",
            accession="MK312485.1",
            accession_url="https://www.ncbi.nlm.nih.gov/nucleotide/MK312485.1?report=genbank&log$=nucltop&blast_rank=4&RID=PYMEHZGA013",
        ),
        BlastNCBIResults(
            description="Bacillus toyonensis strain G9L2 16S ribosomal RNA gene partial sequence",
            description_url="https://blast.ncbi.nlm.nih.gov/Blast.cgi#alnHdr_1362598029",
            scientific_name="Bacillus toyonensis",
            scientific_name_url="https://www.ncbi.nlm.nih.gov/Taxonomy/Browser/wwwtax.cgi?id=155322",
            max_score="1842",
            total_score="1842",
            query_cover="80%",
            e_value="0.0",
            per_indentity="96.44%",
            accession_len="1470",
            accession="MH071323.1",
            accession_url="https://www.ncbi.nlm.nih.gov/nucleotide/MH071323.1?report=genbank&log$=nucltop&blast_rank=5&RID=PYMEHZGA013",
        ),
    ]

    # save_results_in_word("test", species, alignments)
    save_alignments_to_notes(
        r"C:\Users\alber\Desktop\Automatización_secuencias_TFG_Clau/Ali/",
        "test",
        alignments,
    )
